
import docx
from docx.shared import Pt
import os

def create_sample_invoice():
    doc = docx.Document()
    
    # Header
    doc.add_heading('INVOICE', 0)
    
    # Meta info
    p = doc.add_paragraph()
    p.add_run('Invoice Number: ').bold = True
    p.add_run('INV-2024-DOCX-001\n')
    p.add_run('Date: ').bold = True
    p.add_run('2024-12-09\n')
    
    # Vendor
    p.add_run('Vendor: ').bold = True
    p.add_run('Acme Corp\n')
    
    # PO
    p.add_run('PO Number: ').bold = True
    p.add_run('PO-2024-001')
    
    doc.add_paragraph('--------------------------------------------------')
    
    # Line items
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Description'
    hdr_cells[1].text = 'Quantity'
    hdr_cells[2].text = 'Unit Price'
    hdr_cells[3].text = 'Total'
    
    items = [
        ('Widget A', 10, 100.00, 1000.00),
        ('Widget B', 5, 200.00, 1000.00),
        ('Service Fee', 1, 500.00, 500.00)
    ]
    
    for desc, qty, price, total in items:
        row_cells = table.add_row().cells
        row_cells[0].text = desc
        row_cells[1].text = str(qty)
        row_cells[2].text = f"${price:.2f}"
        row_cells[3].text = f"${total:.2f}"
    
    doc.add_paragraph('\n')
    
    # Total
    p = doc.add_paragraph()
    p.add_run('Total Amount: ').bold = True
    p.add_run('$2,500.00')
    
    filename = "test_invoice.docx"
    doc.save(filename)
    print(f"Created {filename}")

if __name__ == "__main__":
    try:
        create_sample_invoice()
    except ImportError:
        print("python-docx not installed")
